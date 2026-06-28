"""
Markdown转DOCX转换脚本
需要先安装: pip install python-docx markdown
"""
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import os

def md_to_docx(md_file, docx_file):
    """将Markdown文件转换为DOCX文件"""
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为HTML
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 处理HTML内容
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # 标题处理
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        # 表格处理（简化）
        elif line.startswith('|'):
            continue  # 表格处理较复杂，暂时跳过
        elif line.startswith('```'):
            continue  # 代码块处理较复杂，暂时跳过
        # 无序列表
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # 普通文本
        else:
            # 移除markdown格式符号
            clean_line = line.replace('**', '').replace('*', '').replace('`', '')
            if clean_line:
                doc.add_paragraph(clean_line)
    
    # 保存DOCX
    doc.save(docx_file)
    print(f"已转换为: {docx_file}")

if __name__ == "__main__":
    md_file = "课程设计报告.md"
    docx_file = "课程设计报告.docx"
    
    if os.path.exists(md_file):
        md_to_docx(md_file, docx_file)
    else:
        print(f"找不到文件: {md_file}")
